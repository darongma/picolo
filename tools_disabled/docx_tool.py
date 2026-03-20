"""
Read, create, and edit Word (.docx) documents.
"""
import docx

# Read tool
tool_spec_read = {
    "name": "docx_read",
    "description": "Extract all text from a Word document (.docx). Returns plain text with paragraph breaks.",
    "parameters": {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the .docx file."
            }
        },
        "required": ["file_path"]
    }
}

def run_read(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs) if paragraphs else "(document has no text)"
    except Exception as e:
        return f"Error: {e}"

# Create tool
tool_spec_create = {
    "name": "docx_create",
    "description": "Create a new Word document (.docx) with the given paragraphs. Each string in the list becomes a paragraph.",
    "parameters": {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path where the .docx file will be saved."
            },
            "paragraphs": {
                "type": "array",
                "items": {"type": "string"},
                "description": "List of paragraphs (strings) to include in the document."
            }
        },
        "required": ["file_path", "paragraphs"]
    }
}

def run_create(file_path: str, paragraphs: list) -> str:
    try:
        doc = docx.Document()
        for text in paragraphs:
            if text.strip():
                doc.add_paragraph(text)
        doc.save(file_path)
        return f"Created Word document with {len(paragraphs)} paragraphs at {file_path}"
    except Exception as e:
        return f"Error: {e}"

# Edit tool
tool_spec_edit = {
    "name": "docx_edit",
    "description": "Edit an existing Word document. Supports append (add paragraphs to the end) or replace (overwrite all paragraphs).",
    "parameters": {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the .docx file to edit."
            },
            "paragraphs": {
                "type": "array",
                "items": {"type": "string"},
                "description": "List of paragraphs to add or replace."
            },
            "append": {
                "type": "boolean",
                "description": "If true, append paragraphs; if false (default), replace all existing content."
            }
        },
        "required": ["file_path", "paragraphs"]
    }
}

def run_edit(file_path: str, paragraphs: list, append: bool = False) -> str:
    try:
        if append:
            doc = docx.Document(file_path)
            for text in paragraphs:
                if text.strip():
                    doc.add_paragraph(text)
            doc.save(file_path)
            return f"Appended {len(paragraphs)} paragraphs to {file_path}"
        else:
            doc = docx.Document()
            for text in paragraphs:
                if text.strip():
                    doc.add_paragraph(text)
            doc.save(file_path)
            return f"Replaced content of {file_path} with {len(paragraphs)} paragraphs"
    except Exception as e:
        return f"Error: {e}"

# Register all tools
tool_specs = [tool_spec_read, tool_spec_create, tool_spec_edit]

tools = {
    "docx_read": run_read,
    "docx_create": run_create,
    "docx_edit": run_edit
}
